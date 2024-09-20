import openai
import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from flask import Flask, request, jsonify, render_template, session, send_file
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import json
from docx import Document


# Flask
app = Flask(__name__)

# secret_key
app.secret_key = os.urandom(24)  

# API key
openai.api_key = 'sk-8dRvpP36C7OaiO2EHeD13SQ7X5Qndi3V7xw9p54m1ET3BlbkFJkISBmIbVbBOOZXlEUGhPI3Ytck_Y9CTmjUISJHDZ8A'

# upload folder
UPLOAD_FOLDER = './input_pdfs'
OUTPUT_FOLDER = './output_texts'
PDF_OUTPUT_FOLDER = './output_pdfs'

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)
if not os.path.exists(PDF_OUTPUT_FOLDER):
    os.makedirs(PDF_OUTPUT_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# 1. PyMuPDF
def pdf_to_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""

    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text += page.get_text("text")  

        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            image = Image.open(io.BytesIO(image_bytes))
            text += pytesseract.image_to_string(image)

    return text

# 2. refine text
def refine_text(text):
    refined_text = text.replace('\n', ' ').strip()
    return refined_text

def generate_flashcards(text):
    prompt = f"""Generate concise and clear flashcards from the following content:
    
{text}

Make sure the output is in the following JSON format:
[
    {{
        "question": "Your question here",
        "answer": "The corresponding answer here"
    }},
    {{
        "question": "Another question",
        "answer": "Another answer"
    }}
]
no ```json or any other text in the beginning or end.
Ensure the output is valid JSON."""
    
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an assistant that helps generate flashcards for study and revision."},
            {"role": "user", "content": prompt}
        ]
    )
    try:
        return json.loads(response['choices'][0]['message']['content'])
    except json.JSONDecodeError:
        return {"error": "Failed to generate flashcards in correct format"}

def generate_cheatsheet(text):
    prompt = f"Generate a concise cheatsheet summarizing the key points from the following content:Focus on main concepts, definitions, formulas, and critical details. Do not return the response in JSON or code format."
    
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an assistant that helps generate concise cheatsheets for study and quick reference."},
            {"role": "user", "content": prompt}
        ]
    )
    return response['choices'][0]['message']['content']

def generate_quiz(text):
    prompt = f"""Generate 5 multiple-choice questions based on the following content:

{text}

Make sure the output is in the following JSON format:
[
    {{
        "question": "What is the question?",
        "choices": [
            "Option A",
            "Option B",
            "Option C",
            "Option D"
        ],
        "correct_answer": "Option B"
    }},
    {{
        "question": "Another question here",
        "choices": [
            "Option A",
            "Option B",
            "Option C",
            "Option D"
        ],
        "correct_answer": "Option D"
    }}
]
no ```json or any other text in the beginning or end.
Ensure the output is valid JSON."""
    
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an assistant that helps generate multiple-choice quiz questions for study and testing purposes."},
            {"role": "user", "content": prompt}
        ]
    )
    try:
        return json.loads(response['choices'][0]['message']['content'])
    except json.JSONDecodeError:
        return {"error": "Failed to generate quiz in correct format"}

def save_cheatsheet_as_word(cheatsheet_content, filename):
    
    doc = Document()

    doc.add_heading('Cheatsheet', level=1)

    for line in cheatsheet_content.split('\n'):
        if line.strip():  
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')  

    # save Word
    doc_path = os.path.join(PDF_OUTPUT_FOLDER, filename)  
    doc.save(doc_path)

    return doc_path

# 4. deal with PDF
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    # PDF to text
    text_content = pdf_to_text(file_path)
    refined_text = refine_text(text_content)

    flashcards = generate_flashcards(refined_text)
    cheatsheet = generate_cheatsheet(refined_text)
    quiz = generate_quiz(refined_text)

    # Save the cheatsheet to Word and provide the download link
    word_filename = f"{os.path.splitext(filename)[0]}_cheatsheet.docx"
    word_path = save_cheatsheet_as_word(cheatsheet, word_filename)

    # Return the response with Word download URL instead of PDF
    return jsonify({
        "flashcards": flashcards,      # Flashcards are already in JSON format
        "cheatsheet": cheatsheet,      # Cheatsheet is a string format
        "quiz": quiz,                  # Quiz data is in JSON format
        "pdf_download_url": f"/download/{word_filename}"  # Link to download the Word doc
    })


@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(PDF_OUTPUT_FOLDER, filename)
    return send_file(file_path, as_attachment=True)

@app.route('/')
def index():
    return render_template('tutor.html')

if __name__ == '__main__':
    app.run(debug=True)